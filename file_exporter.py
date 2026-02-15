'''from fpdf import FPDF
from docx import Document

def save_as_pdf(text, filename="output.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, text)
    pdf.output(filename)
    return filename

def save_as_docx(text, filename="output.docx"):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)
    return filename
'''

from fpdf import FPDF
from docx import Document

def save_as_pdf(text, filename="output.pdf"):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.multi_cell(0, 10, text)
    pdf.output(filename)
    return filename

def save_as_docx(text, filename="output.docx"):
    doc = Document()
    doc.add_paragraph(text)
    doc.save(filename)
    return filename
